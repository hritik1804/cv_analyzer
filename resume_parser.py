# resume_parser.py
import io
import json
import re
from typing import Dict, List, Set, Tuple

from langdetect import detect
from janome.tokenizer import Tokenizer
from rapidfuzz import fuzz
from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text

JP_TOKENIZER = Tokenizer()

def load_skills_dict(path: str) -> Dict[str, List[str]]:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    with io.BytesIO(file_bytes) as bio:
        return pdf_extract_text(bio) or ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    with io.BytesIO(file_bytes) as bio:
        doc = Document(bio)
    texts = []
    for p in doc.paragraphs:
        if p.text:
            texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text)
    return "\n".join(texts)

def detect_lang(text: str) -> str:
    try:
        lang = detect(text)
    except Exception:
        lang = "en"
    return "ja" if "ja" in lang else "en"

def tokenize_en(text: str) -> List[str]:
    # simple baseline tokenization for skills
    # keep words, alphanumerics, and tech tokens
    return re.findall(r"[A-Za-z0-9\+#\.]{2,}", text)

def tokenize_ja(text: str) -> List[str]:
    # extract content words (nouns) and common skill POS
    tokens = []
    for t in JP_TOKENIZER.tokenize(text):
        pos = t.part_of_speech.split(",")[0]
        surface = t.surface.strip()
        if pos in ("名詞",) and surface:
            tokens.append(surface)
    return tokens

def normalize(token: str) -> str:
    return token.strip().lower()

def skills_from_tokens(tokens: List[str], skills_dict: Dict[str, List[str]], fuzzy_threshold: int = 92) -> Set[str]:
    found = set()
    norm_tokens = [normalize(t) for t in tokens]
    for canonical, variants in skills_dict.items():
        canon_norm = normalize(canonical)
        all_variants = {normalize(v) for v in variants + [canonical]}
        # exact hit
        if any(v in norm_tokens for v in all_variants):
            found.add(canonical)
            continue
        # fuzzy hit (good for minor spelling/katakana variations)
        for tok in norm_tokens:
            for v in all_variants:
                if fuzz.ratio(tok, v) >= fuzzy_threshold:
                    found.add(canonical)
                    break
            if canonical in found:
                break
    return found

def parse_resume(file_bytes: bytes, filename: str, skills_dict: Dict[str, List[str]]) -> Tuple[str, str, List[str], Set[str]]:
    if filename.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename.lower().endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file type: use PDF or DOCX")

    lang = detect_lang(text)
    tokens = tokenize_ja(text) if lang == "ja" else tokenize_en(text)
    skill_hits = skills_from_tokens(tokens, skills_dict)
    return lang, text, tokens, skill_hits
